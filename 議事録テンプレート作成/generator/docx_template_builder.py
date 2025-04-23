from docx import Document
from docx.shared import Pt, Inches
from collections import Counter  # ← 忘れずに！

def build_template(common_structure, output_path, style_info=None):
    doc = Document()

    # 書式情報がある場合は設定する
    if style_info:
        section = doc.sections[0]
        section.page_width = style_info["page_width"]
        section.page_height = style_info["page_height"]

        if style_info["orientation"] == "Landscape":
            section.orientation = 1
        else:
            section.orientation = 0

        section.top_margin = style_info["margin_top"]
        section.bottom_margin = style_info["margin_bottom"]
        section.left_margin = style_info["margin_left"]
        section.right_margin = style_info["margin_right"]

    for tag, content in common_structure:
        if tag == "heading":
            para = doc.add_heading("＜見出し＞", level=1)
        else:
            para = doc.add_paragraph("＜内容＞")
        
        if style_info and style_info["fonts"]:
            run = para.runs[0]
            run.font.name = style_info["fonts"][0]
        if style_info and style_info["font_sizes"]:
            run = para.runs[0]
            run.font.size = Pt(style_info["font_sizes"][0])

    doc.save(output_path)  # ← これが build_template の最後

def extract_common_structure(structure_lists):
    tag_sequences = [tuple(tag for tag, _ in struct) for struct in structure_lists]
    tag_counts = Counter(tag_sequences)

    if not tag_counts:
        return []

    most_common_tags = tag_counts.most_common(1)[0][0]
    return [(tag, "（内容を記入）") for tag in most_common_tags]
