from docx import Document

def parse_docx_structure(file_path):
    doc = Document(file_path)
    structure = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # ヘッダー判定の拡張
        if para.style.name.lower().startswith("heading") or (len(text) < 20 and text.endswith(("：", "：", "・", "）"))):
            structure.append(("heading", text))
        else:
            structure.append(("paragraph", text))

    return structure


def get_style_info(file_path):
    doc = Document(file_path)
    section = doc.sections[0]

    style_info = {
        "page_width": section.page_width,
        "page_height": section.page_height,
        "orientation": "Landscape" if section.orientation == 1 else "Portrait",
        "margin_top": section.top_margin,
        "margin_bottom": section.bottom_margin,
        "margin_left": section.left_margin,
        "margin_right": section.right_margin,
        "fonts": [],
    }

    fonts = set()
    font_sizes = set()

    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name:
                fonts.add(run.font.name)
            if run.font.size:
                font_sizes.add(run.font.size.pt)

    style_info["fonts"] = list(fonts)
    style_info["font_sizes"] = list(font_sizes)

    return style_info



